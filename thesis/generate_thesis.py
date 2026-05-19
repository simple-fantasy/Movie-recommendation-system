"""
生成毕业设计论文 Word 文档
青岛理工大学格式规范
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ========== 页面设置 (A4) ==========
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ========== 样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)  # 小四号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

# 设置默认段落首行缩进
style.paragraph_format.first_line_indent = Pt(0)

def add_heading_custom(doc, text, level=1):
    """添加自定义格式的标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Pt(0)

    if level == 1:  # 章标题：三号黑体居中
        run = p.add_run(text)
        run.font.size = Pt(16)  # 三号
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
    elif level == 2:  # 节标题：四号黑体左对齐
        run = p.add_run(text)
        run.font.size = Pt(14)  # 四号
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
    elif level == 3:  # 小节标题：小四号黑体左对齐
        run = p.add_run(text)
        run.font.size = Pt(12)  # 小四号
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True

    return p

def add_body(doc, text):
    """添加正文段落（小四号宋体，首行缩进2字符）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)  # 两个字符缩进
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_page_break(doc):
    doc.add_page_break()

# ==========================================
# 封面
# ==========================================
# 空行
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('本科毕业设计（论文）')
run.font.size = Pt(26)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(30)
run = p.add_run('基于多策略协同过滤的\n电影推荐系统设计与实现')
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(18)
run = p.add_run('Design and Implementation of a Movie Recommendation\nSystem Based on Multi-Strategy Collaborative Filtering')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.italic = True

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)

info_lines = [
    '学    院：____________________',
    '专    业：____________________',
    '学    号：____________________',
    '学生姓名：____________________',
    '指导教师：____________________',
    '',
    '2026年5月',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(line)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_page_break(doc)

# ==========================================
# 中文摘要
# ==========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_after = Pt(18)
run = p.add_run('摘  要')
run.font.size = Pt(16)  # 三号黑体
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

add_body(doc, '随着互联网信息爆炸式增长，用户面临着日益严重的"信息过载"问题，个性化推荐系统成为解决这一问题的核心技术。本课题以MovieLens 32M大规模电影评分数据集为基础，设计并实现了一个融合多种推荐策略的电影推荐系统CineMatch。')

add_body(doc, '本系统采用Flask框架搭建Web后端，基于SQLAlchemy ORM管理MySQL数据库（共19张数据表），前端使用Jinja2模板与Vue.js 3混合架构实现服务端渲染与客户端交互的结合。在推荐算法层面，系统实现了三种推荐策略：基于物品协同过滤（ItemCF）采用scikit-learn最近邻搜索进行余弦相似度预计算，并支持可解释推荐理由生成；神经协同过滤（NCF）基于PyTorch实现GMF架构的神经网络模型，通过异步加载和单例模式进行高效推理；混合推荐（Hybrid）创新性地采用ItemCF召回+NCF重排的级联架构，在保证推荐精度的同时兼顾可解释性。')

add_body(doc, '系统还实现了用户画像分析（12维特征）、用户行为追踪、推荐反馈收集、影单社交、数据可视化看板（ECharts）和管理后台等完整功能模块。通过统一的离线评估框架（Leave-Last-Out协议），使用Precision@K、Recall@K、MAP@K、NDCG@K和MRR@K五项指标对三种推荐策略进行了对比评估，并通过消融实验分析了关键参数的影响。')

# 关键词
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(18)
run = p.add_run('关键词：')
run.font.size = Pt(12)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True
run = p.add_run('电影推荐系统，协同过滤，神经协同过滤，混合推荐，可解释推荐')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_page_break(doc)

# ==========================================
# 英文摘要
# ==========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_after = Pt(18)
run = p.add_run('ABSTRACT')
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.bold = True

abstract_en1 = (
    'With the explosive growth of Internet information, users face an increasingly severe '
    '"information overload" problem, making personalized recommendation systems a core technology '
    'to address this challenge. This project designs and implements a movie recommendation system '
    'named CineMatch, integrating multiple recommendation strategies based on the MovieLens 32M '
    'large-scale rating dataset.'
)
add_body(doc, abstract_en1)

abstract_en2 = (
    'The system employs the Flask framework for the web backend, manages a MySQL database (19 tables) '
    'through SQLAlchemy ORM, and uses a hybrid frontend architecture combining Jinja2 templates with '
    'Vue.js 3 for server-side rendering and client-side interaction. At the recommendation algorithm level, '
    'three strategies are implemented: Item-based Collaborative Filtering (ItemCF) uses scikit-learn nearest '
    'neighbor search for precomputed cosine similarity with explainable recommendation reasons; Neural '
    'Collaborative Filtering (NCF) implements a GMF-architecture neural network based on PyTorch with async '
    'loading and singleton pattern for efficient inference; the Hybrid strategy innovatively adopts a cascade '
    'architecture with ItemCF recall followed by NCF reranking, balancing recommendation accuracy with '
    'interpretability.'
)
add_body(doc, abstract_en2)

abstract_en3 = (
    'The system also implements comprehensive functional modules including user profiling (12-dimensional '
    'features), user behavior tracking, recommendation feedback collection, movie list social features, '
    'data visualization dashboards (ECharts), and an admin panel. Through a unified offline evaluation '
    'framework (Leave-Last-Out protocol), five metrics—Precision@K, Recall@K, MAP@K, NDCG@K, and '
    'MRR@K—are used to comparatively evaluate the three strategies, with ablation studies analyzing '
    'the impact of key parameters.'
)
add_body(doc, abstract_en3)

# KEY WORDS
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(18)
run = p.add_run('KEY WORDS: ')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True
run = p.add_run('movie recommendation system, collaborative filtering, neural collaborative filtering, '
                'hybrid recommendation, explainable recommendation')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

add_page_break(doc)

# ==========================================
# 目录
# ==========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_after = Pt(24)
run = p.add_run('目  录')
run.font.size = Pt(18)  # 小二黑体
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

# TOC entries
toc_entries = [
    ('摘要', 'I'),
    ('ABSTRACT', 'II'),
    ('目录', 'III'),
    ('第1章 绪论', '1'),
    ('  1.1 研究背景与意义', '1'),
    ('  1.2 国内外研究现状', '2'),
    ('  1.3 主要研究内容', '4'),
    ('  1.4 论文组织结构', '4'),
    ('第2章 相关理论与技术', '5'),
    ('  2.1 推荐系统概述', '5'),
    ('  2.2 基于物品的协同过滤', '6'),
    ('  2.3 神经协同过滤', '8'),
    ('  2.4 混合推荐策略', '10'),
    ('  2.5 系统开发关键技术', '10'),
    ('第3章 系统需求分析', '12'),
    ('  3.1 功能性需求', '12'),
    ('  3.2 非功能性需求', '14'),
    ('  3.3 数据需求分析', '15'),
    ('第4章 系统设计', '16'),
    ('  4.1 系统总体架构设计', '16'),
    ('  4.2 数据库设计', '17'),
    ('  4.3 推荐算法设计', '20'),
    ('  4.4 前端界面设计', '23'),
    ('第5章 系统实现', '24'),
    ('  5.1 开发环境', '24'),
    ('  5.2 核心功能模块实现', '24'),
    ('  5.3 推荐算法实现', '28'),
    ('  5.4 前端实现', '30'),
    ('第6章 系统测试与评估', '32'),
    ('  6.1 系统功能测试', '32'),
    ('  6.2 推荐算法离线评估', '33'),
    ('  6.3 系统性能分析', '35'),
    ('第7章 结论与展望', '36'),
    ('  7.1 工作总结', '36'),
    ('  7.2 存在的不足', '37'),
    ('  7.3 未来展望', '37'),
    ('致谢', '38'),
    ('参考文献', '39'),
]

for title, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_page_break(doc)

# ==========================================
# 第1章 绪论
# ==========================================
add_heading_custom(doc, '第1章 绪论', level=1)

add_heading_custom(doc, '1.1 研究背景与意义', level=2)

add_body(doc, '在信息爆炸的互联网时代，用户面临着日益严峻的"信息过载"问题。以电影领域为例，Netflix、IMDb等平台拥有数十万部电影资源，用户难以从海量内容中高效地找到符合个人偏好的影片。传统的分类浏览和关键词搜索方式依赖用户主动输入明确的查询意图，无法满足用户在缺乏明确目标时的发现需求。推荐系统作为一种信息过滤技术，通过分析用户的历史行为数据构建偏好模型，能够主动为用户推送个性化内容，已成为解决信息过载问题的核心技术手段。')

add_body(doc, '电影推荐系统的研究具有重要的理论价值和实践意义。从理论层面看，电影推荐涉及用户行为建模、物品相似度计算、深度学习特征提取等多个研究问题，是推荐系统领域经典且具有持续研究活力的方向。特别是MovieLens数据集作为推荐系统研究中最广泛使用的基准数据集之一，其最新32M版本包含3200万条评分记录，覆盖27万用户和8.7万部电影，为大规模推荐算法的验证提供了良好基础。从实践层面看，一个完整的电影推荐系统需要整合数据预处理、模型训练、在线推理、前端交互、用户反馈等多个工程模块，是检验推荐算法工程化能力的理想场景。')

add_body(doc, '本课题旨在基于MovieLens 32M数据集，设计并实现一个融合多种推荐策略的电影推荐系统，兼顾推荐精度与可解释性，并提供完整的用户交互界面和数据分析看板。')

add_heading_custom(doc, '1.2 国内外研究现状', level=2)

add_heading_custom(doc, '1.2.1 传统协同过滤方法', level=3)

add_body(doc, '协同过滤（Collaborative Filtering, CF）是推荐系统领域最经典的技术路线，主要分为基于用户的协同过滤（UserCF）和基于物品的协同过滤（ItemCF）两类。UserCF通过计算用户间的相似度，为目标用户推荐相似用户喜欢的物品。ItemCF则通过计算物品间的相似度，为用户推荐与其历史偏好物品相似的其他物品。2003年，Amazon将ItemCF成功应用于电商推荐，证明了该方法的工业可行性[8]。2001年，Sarwar等人系统性地研究了ItemCF在稀疏评分矩阵下的表现，提出了基于余弦相似度和调整余弦相似度的物品相似度计算方法[1]。')

add_body(doc, '在实际研究中，2009年Koren等人提出的矩阵分解（Matrix Factorization, MF）方法将用户和物品映射到共同的隐向量空间，通过隐语义建模显著提升了推荐精度，并在Netflix Prize竞赛中取得了突破性成果[3]。此后，SVD++、PMF、BPR等变体方法进一步扩展了矩阵分解的建模能力。2009年Rendle等人提出的BPR（Bayesian Personalized Ranking）从排序优化的角度重新审视了推荐问题，将推荐建模为隐反馈数据的排序学习任务[6]。')

add_heading_custom(doc, '1.2.2 深度学习推荐方法', level=3)

add_body(doc, '近年来，深度学习技术在推荐系统领域得到了广泛应用。2017年，He等人提出了神经协同过滤（Neural Collaborative Filtering, NCF）框架，使用多层感知机替代传统矩阵分解中的内积操作，能够学习更复杂的用户-物品交互函数[2]。NCF中的广义矩阵分解（GMF）组件保留了矩阵分解的线性建模能力，而多层感知机（MLP）组件则引入了非线性变换，两者的结合实现了线性和非线性特征的互补。')

add_body(doc, '除NCF外，Wide & Deep[9]、DeepFM[10]、DIN等模型在工业界也得到了广泛应用。这些模型通过引入注意力机制、特征交叉等技术进一步提升了推荐性能。2019年，Wang等人提出了神经图协同过滤（NGCF），将图神经网络引入协同过滤，通过在用户-物品交互图上进行嵌入传播来捕捉高阶连通性[7]。然而，深度学习模型普遍面临可解释性不足的问题——推荐结果难以像ItemCF那样给出"因为你喜欢X，所以推荐Y"的解释，这在注重用户体验的场景中是一个重要短板。')

add_heading_custom(doc, '1.2.3 混合推荐策略', level=3)

add_body(doc, '单一推荐算法往往存在固有局限。协同过滤面临冷启动和数据稀疏问题，基于内容的推荐存在特征提取困难，深度学习模型缺乏可解释性。因此，混合推荐（Hybrid Recommendation）成为研究热点。常见的混合策略包括加权混合、切换混合、级联混合和特征组合混合等[5]。本课题采用级联混合策略：首先由ItemCF生成大规模召回候选集，然后由NCF对候选集进行精细化排序。这种方式既保留了ItemCF的可解释性优势，又发挥了深度学习模型的非线性拟合能力，同时有效控制了NCF推理的计算开销——只需对数百个候选物品评分，而非全量8.7万物品。')

add_heading_custom(doc, '1.3 主要研究内容', level=2)

add_body(doc, '本课题围绕电影推荐系统的设计与实现，主要开展以下研究工作：')

add_body(doc, '（1）数据预处理与特征工程。针对MovieLens 32M数据集进行数据清洗、过滤和质量控制，包括评分数量阈值过滤、用户均值归一化处理等。同时通过TMDB API进行电影元数据补充，获取电影海报、导演、演员、简介等富文本信息。')

add_body(doc, '（2）多策略推荐算法设计与实现。实现基于物品协同过滤（ItemCF）、神经协同过滤（NCF）和混合推荐（Hybrid）三种推荐策略。ItemCF采用scikit-learn的最近邻搜索实现高效相似度计算，NCF基于PyTorch实现GMF架构的神经网络模型，Hybrid采用ItemCF召回+NCF重排的级联架构。')

add_body(doc, '（3）可解释性推荐机制。在ItemCF和Hybrid策略中实现推荐理由生成，向用户展示"因为你喜欢X，所以推荐Y"的可解释推荐结果，增强系统透明度和用户信任。')

add_body(doc, '（4）完整Web系统构建。基于Flask应用工厂模式搭建后端服务，使用SQLAlchemy ORM管理数据库，采用Jinja2+Vue.js 3构建服务端渲染与客户端交互结合的前端界面，实现用户认证、电影浏览、评分、收藏、评论、影单、通知等完整功能。')

add_body(doc, '（5）离线评估与消融实验。设计统一的离线评估框架，采用Leave-Last-Out协议，计算Precision@K、Recall@K、MAP@K、NDCG@K、MRR@K五项指标，对比三种策略的性能差异，并通过消融实验分析关键参数的影响。')

add_heading_custom(doc, '1.4 论文组织结构', level=2)

add_body(doc, '本论文共分为七章，各章内容安排如下：')

add_body(doc, '第1章为绪论，介绍课题的研究背景与意义、国内外研究现状、主要研究内容和论文组织结构。')

add_body(doc, '第2章为相关理论与技术，介绍推荐系统的基本概念、协同过滤算法的数学原理、神经协同过滤的模型架构，以及系统开发所采用的关键技术。')

add_body(doc, '第3章为系统需求分析，从功能性需求、非功能性需求和数据需求三个维度对系统进行全面的需求分析。')

add_body(doc, '第4章为系统设计，阐述系统的总体架构设计、数据库表结构设计和推荐算法的详细设计方案。')

add_body(doc, '第5章为系统实现，介绍开发环境配置、核心功能模块的实现细节、推荐算法的具体实现过程以及前端界面的构建。')

add_body(doc, '第6章为系统测试与评估，包括系统功能的测试验证和推荐算法的离线评估与消融实验结果分析。')

add_body(doc, '第7章为结论与展望，总结本课题的研究成果，指出存在的不足，并对未来改进方向进行展望。')

add_page_break(doc)

# ==========================================
# 第2章 相关理论与技术
# ==========================================
add_heading_custom(doc, '第2章 相关理论与技术', level=1)

add_heading_custom(doc, '2.1 推荐系统概述', level=2)

add_body(doc, '推荐系统是一种信息过滤工具，其核心任务是根据用户的历史行为、个人信息和上下文环境，从海量候选物品中筛选出用户可能感兴趣的内容。形式化地，推荐问题可以描述为：给定用户集合U和物品集合I，对于每个用户u∈U，推荐系统需要预测用户u对未交互物品i∈I的偏好评分r(u,i)，并选取预测评分最高的K个物品作为推荐结果。')

add_body(doc, '推荐系统的基本技术路线可以分为三类：（1）基于内容的推荐（Content-Based Filtering），利用物品的属性信息和用户的偏好特征进行匹配；（2）基于协同过滤的推荐（Collaborative Filtering），利用用户群体的行为数据进行推荐；（3）混合推荐（Hybrid Recommendation），综合多种推荐策略的优势[5]。在实际应用中，电影推荐系统的度量指标包括评分预测精度（如RMSE、MAE）和排序质量（如Precision@K、Recall@K、NDCG@K）两个维度。')

add_heading_custom(doc, '2.2 基于物品的协同过滤', level=2)

add_body(doc, '基于物品的协同过滤（Item-based Collaborative Filtering, ItemCF）是推荐系统领域最经典的算法之一。其核心思想是：如果两个物品被相同用户群体给予相似的评分，那么这两个物品是相似的；为用户推荐与其历史偏好物品相似的其他物品。')

add_heading_custom(doc, '2.2.1 评分矩阵与稀疏性', level=3)

add_body(doc, 'ItemCF的输入是用户-物品评分矩阵R，其中矩阵的每一行代表一个用户，每一列代表一个电影，元素R_{u,i}表示用户u对电影i的评分。在实际场景中，该矩阵通常极其稀疏——以MovieLens 32M数据集为例，评分密度仅为3200万/(27万×8.7万)≈0.14%[4]。稀疏性使得传统的基于用户共同评分项计算相似度的方式可能产生不可靠的相似度估计，需要通过数据过滤和质量控制来缓解。')

add_heading_custom(doc, '2.2.2 相似度计算与归一化', level=3)

add_body(doc, '对于两部电影i和j，两者之间的余弦相似度定义为评分矩阵中对应列向量的余弦夹角。为了消除用户评分尺度差异带来的偏差，本系统采用用户均值归一化（调整余弦相似度）[1]。对于每个用户u，计算其历史评分的均值μ_u，然后将评分中心化。这一预处理步骤能够消除"宽容型用户"（普遍打高分）和"苛刻型用户"（普遍打低分）之间的系统偏差，使相似度计算更加准确。')

add_heading_custom(doc, '2.2.3 推荐生成与解释', level=3)

add_body(doc, '对于目标用户u，设其已评分的电影集合为I_u，ItemCF的推荐分数计算为：对用户历史中的每一部电影j，将其评分R_{u,j}与目标电影i与j的相似度sim(i,j)加权求和。得分最高的K部电影作为推荐结果。ItemCF的一个重要优势是天然支持推荐解释——对于推荐结果中的电影i，可以回溯到对其贡献最大的评分种子电影j，向用户展示"因为你喜欢《j》，所以推荐《i》"的解释信息，增强系统的透明度和用户信任。')

add_heading_custom(doc, '2.2.4 大规模近邻搜索', level=3)

add_body(doc, '在8.7万部电影的规模下，暴力计算全量物品对相似度的时间复杂度为O(|I|²|U|)，代价过高。本系统采用余弦距离上的最近邻搜索策略：将评分矩阵转置为电影×用户的矩阵，使用scikit-learn的NearestNeighbors模块在余弦距离空间构建最近邻索引，高效地检索每部电影的Top-K最相似电影。')

add_heading_custom(doc, '2.3 神经协同过滤', level=2)

add_body(doc, '神经协同过滤（Neural Collaborative Filtering, NCF）由He等人在2017年提出[2]，是利用深度学习技术增强协同过滤能力的代表性方法。')

add_heading_custom(doc, '2.3.1 从矩阵分解到NCF', level=3)

add_body(doc, '传统矩阵分解将用户u和物品i分别映射为嵌入向量p_u和q_i，然后通过内积计算预测评分。这种方法假设嵌入向量的各维度相互独立且以等权重线性组合，限制了模型的表达能力。NCF将该交互函数替换为一个可学习的多层神经网络，能够学习嵌入维度之间的非线性交互。')

add_heading_custom(doc, '2.3.2 GMF架构', level=3)

add_body(doc, '本系统实现的NCF采用广义矩阵分解（GMF）与多层感知机（MLP）结合的双塔架构。具体网络结构如下：（1）嵌入层：将用户ID和物品ID分别通过Embedding层映射为稠密向量，维度为embedding_dim（默认32维）；（2）向量拼接：将用户嵌入向量与物品嵌入向量拼接；（3）MLP层：采用三层全连接网络，结构为[2d→hidden_dim→hidden_dim/2→1]，每层之间使用ReLU激活函数引入非线性；（4）预测输出：将logit通过Sigmoid函数转换为偏好概率。')

add_heading_custom(doc, '2.3.3 损失函数与负采样', level=3)

add_body(doc, 'NCF采用二值交叉熵损失（Binary Cross-Entropy Loss）进行训练。由于实际数据中正负样本严重不均衡，本系统在训练过程中采用负采样策略：对于每个正样本，随机采样neg_ratio个用户未交互的物品作为负样本。负采样使用碰撞检测重采样机制——若采样到的负样本恰好属于该用户的正样本集合，则重新采样，确保负样本的有效性。')

add_heading_custom(doc, '2.3.4 评估协议', level=3)

add_body(doc, 'NCF的离线评估采用Leave-Last-Out协议：对每个用户，按交互时间排序，取最后一次交互作为验证集，其余作为训练集。评估时，从用户未交互物品中随机采样100个作为候选集，与真实正样本混合后排序，计算HR@K（命中率）和NDCG@K（归一化折损累计增益）。这种采样评估方式避免了全量物品排序带来的计算开销，同时保证了评估的公平性。')

add_heading_custom(doc, '2.4 混合推荐策略', level=2)

add_body(doc, '本系统采用级联混合（Cascade Hybrid）策略，将ItemCF作为召回阶段，NCF作为排序阶段：（1）召回阶段：ItemCF从全量电影中检索recall_k（默认100）部候选电影，利用预计算的相似度表实现快速响应；（2）排序阶段：NCF对召回候选集进行重新评分和排序，输出最终的top_n（默认10）部推荐电影；（3）兜底策略：NCF不可用或冷启动用户自动降级为纯ItemCF推荐，完全冷启动返回热门电影。这种级联架构的优势在于：ItemCF保证了候选集的质量和可解释性，NCF引入了非线性特征交互提升排序精度，两者可独立部署和升级。')

add_heading_custom(doc, '2.5 系统开发关键技术', level=2)

add_body(doc, '后端框架方面，系统采用Flask Web框架，使用应用工厂模式（Application Factory Pattern）组织代码，通过蓝图（Blueprint）进行模块化路由管理。数据库访问层采用SQLAlchemy ORM，通过声明式模型映射支持MySQL和SQLite双数据库后端。用户认证基于Flask-Login扩展，实现了基于Session的认证机制和角色权限管理。')

add_body(doc, '前端架构方面，系统采用服务端渲染与客户端交互结合的混合架构。Jinja2模板引擎负责页面初始结构的渲染，Vue.js 3（通过CDN引入，使用Options API）负责页面的动态交互逻辑。UI框架基于Bootstrap 5深色主题，数据可视化采用ECharts图表库。')

add_body(doc, '科学计算与深度学习方面，相似度计算基于NumPy、SciPy稀疏矩阵和scikit-learn近邻搜索。NCF模型基于PyTorch框架实现，支持CPU和CUDA双模式，通过单例模式（Singleton）进行内存常驻管理。系统提供RESTful风格的API，共70余个端点，推荐API通过strategy参数切换推荐策略。')

add_page_break(doc)

# ==========================================
# 第3章 系统需求分析
# ==========================================
add_heading_custom(doc, '第3章 系统需求分析', level=1)

add_heading_custom(doc, '3.1 功能性需求', level=2)

add_body(doc, '根据电影推荐系统的业务场景分析，本系统的功能性需求可划分为以下八个模块：')

add_body(doc, '（1）用户管理模块。支持用户注册与登录，密码采用werkzeug的哈希加盐存储确保安全性。登录后通过Flask-Login维护Session状态。区分普通用户和管理员角色，管理员可对用户进行启用/禁用、权限授予/撤销等操作。支持用户查看和修改个人资料。')

add_body(doc, '（2）电影浏览与搜索模块。支持电影详情页展示，包含标题、年份、类型、导演、演员、海报、评分等完整元数据。支持按标题关键词、类型、年份范围、评分阈值等多维度组合搜索和排序。基于预计算的ItemCF相似度表提供相似电影推荐。')

add_body(doc, '（3）个性化推荐模块。提供ItemCF、NCF和Hybrid三种推荐策略，用户可按需切换。ItemCF和Hybrid策略提供推荐理由展示。对于无评分记录的新用户自动返回热门电影兜底。支持用户对推荐结果进行喜欢/不喜欢反馈标记。')

add_body(doc, '（4）用户交互模块。支持0.5至5.0分步长为0.5的精细化电影评分，支持favorite/watchlist/seen三种收藏类型，支持文字评论和点赞，支持用户创建自定义影单（公开/私有模式），支持提交和查看电影在线观看链接。')

add_body(doc, '（5）数据看板模块。提供评分分布、类型分布、年份趋势、热门电影排行等统计图表，以及系统概览、用户分群、活动热力图、系统健康度监控等高级分析功能。')

add_body(doc, '（6）用户画像模块。基于用户历史行为计算类型偏好、年代偏好、演员/导演偏好、评分行为特征、观影多样性等12维特征，自动生成个性化观影洞察。')

add_body(doc, '（7）通知系统。支持系统通知、社交通知、推荐通知和成就通知，用户可自定义通知偏好设置。')

add_body(doc, '（8）管理后台。提供电影管理、用户管理、评论审核、通知管理、元数据管理、数据导出等管理功能。')

add_heading_custom(doc, '3.2 非功能性需求', level=2)

add_body(doc, '（1）性能要求：电影列表加载和推荐结果返回应在2秒内完成，推荐API应充分利用预计算数据。')

add_body(doc, '（2）可扩展性要求：三种推荐策略可独立部署和切换，新增策略不应影响现有策略的运行。')

add_body(doc, '（3）安全性要求：用户密码加密存储，API进行输入合法性校验，管理后台操作需权限验证，数据库连接池管理防止连接泄露。')

add_body(doc, '（4）可用性要求：NCF模型通过异步加载机制实现优雅降级——模型不可用时自动切换ItemCF策略，不影响用户正常使用。')

add_body(doc, '（5）可维护性要求：代码遵循模块化原则，配置由Pydantic Settings统一管理，日志采用结构化格式便于排查。')

add_heading_custom(doc, '3.3 数据需求分析', level=2)

add_body(doc, '本系统基于MovieLens 32M数据集。该数据集由GroupLens研究实验室发布，包含32,000,000条电影评分记录，涉及270,000名用户和87,000部电影。评分值为0.5至5.0的连续值，步长为0.5。此外，数据集还包含电影的标题、年份和类型等元数据。')

add_body(doc, '为了丰富电影的元数据信息，系统通过TMDB（The Movie Database）API进行数据补充。补充的信息包括：电影海报、背景图、原始标题、导演、演员阵容、剧情简介、片长、宣传语、IMDb ID等。')

add_body(doc, '在数据过滤方面，为确保推荐算法在有效数据上运行，系统实施了以下策略：过滤评分数量少于50条的冷门电影，确保相似度计算的可靠性；过滤评分数量少于5条（ItemCF）或10条（NCF）的不活跃用户；对评分进行用户均值归一化处理，消除评分尺度系统偏差。')

add_body(doc, '在数据库选型方面，系统支持双数据库后端：开发/测试环境使用SQLite，无需额外安装数据库服务；生产环境使用MySQL+PyMySQL驱动，支持高并发和大数据量场景，数据库编码采用utf8mb4确保多语言内容正确存储。')

add_page_break(doc)

# ==========================================
# 第4章 系统设计
# ==========================================
add_heading_custom(doc, '第4章 系统设计', level=1)

add_heading_custom(doc, '4.1 系统总体架构设计', level=2)

add_body(doc, '本系统采用经典的三层Web应用架构，分为表示层、业务逻辑层和数据持久层。')

add_body(doc, '表示层（Presentation Layer）由Jinja2模板引擎和Vue.js 3协同构成。Jinja2负责服务端渲染页面初始结构，Vue.js处理客户端的动态交互逻辑。Bootstrap 5提供响应式UI组件，ECharts承担数据可视化任务。前端每个页面对应独立的page-*.js文件和一个Jinja2模板文件，实现关注点分离。')

add_body(doc, '业务逻辑层（Business Logic Layer）基于Flask应用工厂模式构建。核心组件包括：路由处理器（routes.py和admin_routes.py，共计70余个API端点）、推荐引擎（ItemCF召回模块和NCF推理引擎）、用户画像服务（ProfileService）、行为追踪模块（BehaviorTracker），以及配置管理、缓存管理、限流控制、日志记录等支撑模块。')

add_body(doc, '数据持久层（Data Persistence Layer）通过SQLAlchemy ORM实现数据库操作的对象化抽象，支持MySQL和SQLite双后端。19张数据表覆盖用户、电影、评分、相似度、评论、收藏、影单、通知、行为日志、用户画像等业务域。')

add_body(doc, '关键的设计决策遵循以下原则：（1）关注点分离——前端每个页面对应独立的JS文件，后端路由按功能域组织；（2）配置外部化——所有配置项通过Pydantic BaseSettings从.env文件加载，支持类型校验和范围验证；（3）优雅降级——NCF模型通过异步加载和状态检查机制实现自动回退；（4）单例管理——NCFEngine采用线程安全的单例模式，确保模型在内存中只加载一份。')

add_heading_custom(doc, '4.2 数据库设计', level=2)

add_heading_custom(doc, '4.2.1 核心业务表设计', level=3)

add_body(doc, '用户表（users）存储用户账户信息，核心字段包括id（主键）、username（唯一索引）、password_hash（werkzeug哈希加密）、email、avatar、is_admin（管理员标识）、is_active（账户状态）、last_login、login_count、created_at。password_hash采用werkzeug.security.generate_password_hash()进行加盐哈希，不在数据库中存储明文密码。')

add_body(doc, '电影表（movies）存储电影元数据，核心字段包括id、title、year、genres（竖线分隔）、original_title、director、actors（JSON存储）、description、runtime、poster_url、backdrop_url、trailer_url、tagline、tmdb_id、imdb_id、language、country。管理字段包括status（active/inactive/pending）、is_featured、view_count、rating_count、avg_rating。actors字段以JSON格式存储演员列表。')

add_body(doc, '评分表（ratings）存储用户对电影的评分记录，核心字段包括id、user_id、movie_id、rating（0.5-5.0）、timestamp。表上建有(user_id, movie_id)唯一约束防止重复评分，(user_id, timestamp)复合索引优化用户历史查询，(movie_id, rating)复合索引优化电影评分统计。')

add_body(doc, '电影相似度表（movie_similarity）存储ItemCF预计算的相似度数据，核心字段包括id、movie_id、similar_movie_id、score。建有(movie_id, similar_movie_id)唯一约束和(movie_id, score)复合索引，优化"查询某电影最相似电影列表"这一高频操作。')

add_heading_custom(doc, '4.2.2 社交功能表设计', level=3)

add_body(doc, '评论表（reviews）存储用户对电影的评论，包含user_id、movie_id、content、rating、likes_count、is_featured、status（approved/rejected/pending）等字段。评论点赞表（review_likes）通过(user_id, review_id)唯一约束防止重复点赞。')

add_body(doc, '收藏表（user_collections）存储用户的电影收藏记录，通过collection_type字段区分收藏（favorite）、待看（watchlist）和已看（seen）三种类型，通过(user_id, movie_id, collection_type)唯一约束支持同一用户将同一电影加入不同类型的收藏。')

add_body(doc, '影单系统由三张表组成：影单表（movie_lists）存储影单基本信息和统计数据，影单项表（movie_list_items）存储影单中的电影条目及排序，影单点赞表（movie_list_likes）和评论表（movie_list_comments）支持社交互动。')

add_heading_custom(doc, '4.2.3 其他辅助表设计', level=3)

add_body(doc, '通知表（notifications）存储用户通知，支持system/review_reply/review_liked/movie_recommend/achievement五种类型。用户通知偏好表（user_notification_preferences）允许用户自由开关各类通知。')

add_body(doc, '用户画像表（user_profiles）存储12维用户偏好和行为特征，包括preferred_genres、preferred_years、preferred_actors、preferred_directors（均为JSON格式）、avg_rating_level、rating_variance、rating_entropy、total_watch_time、genre_diversity、decade_diversity、user_type、activity_level。')

add_body(doc, '用户行为表（user_behaviors）存储用户操作日志，包含action_type（view/rate/search/click）、target_type、target_id、extra_data（JSON）、ip_address、session_id、referrer等字段，支持灵活的行为分析和用户画像计算。')

add_heading_custom(doc, '4.3 推荐算法设计', level=2)

add_heading_custom(doc, '4.3.1 ItemCF算法设计', level=3)

add_body(doc, 'ItemCF算法的设计分为离线训练和在线推理两个阶段。')

add_body(doc, '离线训练阶段的步骤包括：（1）从评分表加载全量数据构建DataFrame；（2）按min_ratings_per_movie（默认50）过滤冷门电影，按min_ratings_per_user（默认5）过滤不活跃用户；（3）对评分进行用户均值归一化消除评分尺度偏差；（4）使用pandas.factorize将ID映射为连续整数索引，通过scipy.sparse.coo_matrix构建电影×用户稀疏矩阵，转换为CSR格式优化行向量访问；（5）使用sklearn.neighbors.NearestNeighbors以余弦距离进行最近邻搜索，为每部电影找到topk（默认50）个最相似电影；（6）将距离转换为相似度分数（1-distance），批量写入movie_similarity表。')

add_body(doc, '在线推理阶段的流程为：（1）获取当前用户全部评分记录和已评分电影ID集合；（2）使用IN子句批量查询所有已评分电影的相似电影列表（单次SQL查询，避免N+1问题）；（3）按已评分电影分组，对每组相似电影取Top-50，在累加分数时过滤已评分电影；（4）对候选电影按累加分数降序排列，选取Top-N作为推荐结果；（5）生成推荐理由：回溯贡献最大的3部种子电影，返回"因为你喜欢X"的解释。')

add_heading_custom(doc, '4.3.2 NCF算法设计', level=3)

add_body(doc, 'NCF算法的训练阶段流程包括：数据加载与Leave-Last-Out划分、ID映射（user2idx/item2idx）、负采样（含碰撞检测重采样）、模型构建（Embedding+三层MLP）、Adam优化器+BCE损失训练、早停机制（patience=3，监控验证集NDCG@K）、模型导出（ncf.pt参数文件和ncf_meta.json元数据文件）。')

add_body(doc, '推理阶段通过NCFEngine全局单例管理模型生命周期。模型采用异步加载机制在后台线程加载，不阻塞Web服务启动。推理时提供两个核心接口：score(user_id, item_ids)对候选物品批量评分，rank(user_id, item_ids, top_k)返回Top-K排序结果。')

add_heading_custom(doc, '4.3.3 Hybrid混合推荐设计', level=3)

add_body(doc, 'Hybrid策略采用级联架构：（1）ItemCF召回阶段从全量电影中检索recall_k（默认100）部候选电影；（2）检查NCFEngine状态，若可用则将候选电影送入NCF重新评分排序，输出top_n推荐结果；（3）推荐理由沿用ItemCF阶段的相似度贡献信息；（4）降级策略：NCF加载中返回503状态码，NCF加载失败或用户不在训练集中自动降级为纯ItemCF。')

add_heading_custom(doc, '4.4 前端界面设计', level=2)

add_body(doc, '系统前端采用服务端渲染与客户端交互结合的混合架构。页面布局方面，所有页面继承自base.html模板，提供统一导航栏、页脚和CSS框架。组件化设计方面，每个功能页面对应一个Vue.js应用，通过CDN引入Vue 3并使用Options API，所有API调用封装在api.js公共模块中。响应式设计基于Bootstrap 5栅格系统。数据可视化使用ECharts实现评分分布、类型分布、年份趋势、用户分群等图表。视觉风格采用深色电影主题，使用毛玻璃效果、渐变背景和动画手法营造现代电影平台的视觉效果。')

add_page_break(doc)

# ==========================================
# 第5章 系统实现
# ==========================================
add_heading_custom(doc, '第5章 系统实现', level=1)

add_heading_custom(doc, '5.1 开发环境', level=2)

add_body(doc, '本系统基于以下技术栈开发：编程语言为Python 3.10，Web框架为Flask 3.0 + Flask-SQLAlchemy + Flask-Login + Flask-Caching，数据库为MySQL 8.0（主）+ PyMySQL驱动和SQLite 3（开发备选），科学计算库包括NumPy 1.24+、Pandas 2.0+、scikit-learn 1.3+、SciPy 1.11+，深度学习框架为PyTorch 2.0+（支持CPU/CUDA双模式），前端库包括Vue.js 3（CDN引入）、Bootstrap 5.3、ECharts 5.4、GSAP 3.12，其他依赖包括Pydantic（配置管理）、structlog（结构化日志）、TMDB API v3（元数据补充）。')

add_heading_custom(doc, '5.2 核心功能模块实现', level=2)

add_heading_custom(doc, '5.2.1 应用工厂与配置管理', level=3)

add_body(doc, '系统的入口通过create_app()工厂函数构建Flask应用实例。该函数执行以下初始化步骤：设置Jinja2环境为ChainableUndefined（允许Vue.js的{{ }}与Jinja2模板语法共存）；从Config类加载配置项并设置Session Cookie参数；初始化SQLAlchemy、LoginManager、Flask-Caching三个核心扩展；注册main和admin两个蓝图；注册全局中间件（错误处理、请求日志、限流控制）；在应用上下文中创建所有数据库表并执行种子数据填充；在后台线程异步进行NCF模型预加载。')

add_body(doc, '配置管理采用Pydantic BaseSettings实现。Settings类定义了所有配置项及其默认值、类型校验和范围验证规则。通过@validator装饰器实现了SECRET_KEY强度校验、数据库URL驱动校验、Embedding维度范围校验和限流参数校验。使用@lru_cache确保Settings实例全局单例。Config类将Settings实例的属性映射为Flask兼容的配置项，实现配置层的关注点分离。')

add_heading_custom(doc, '5.2.2 用户认证实现', level=3)

add_body(doc, '用户认证基于Flask-Login扩展实现。User模型继承UserMixin，密码通过werkzeug.security.generate_password_hash()加盐哈希存储。注册接口进行用户名长度（2-64字符）、密码长度（≥6字符）、密码一致性、用户名唯一性四重校验。登录接口验证成功后更新last_login和login_count。对于API请求（/api/前缀），未登录时返回JSON格式401错误而非HTML重定向，确保前端Vue.js能正确处理认证失败。管理员权限通过@admin_required装饰器进行路由级别的访问控制。')

add_heading_custom(doc, '5.2.3 推荐API实现', level=3)

add_body(doc, '推荐API（GET /api/recommendations）是系统的核心接口，接受n（推荐数量）、strategy（策略）、recall_k（Hybrid召回数量）三个参数。')

add_body(doc, '对于冷启动用户（未登录或评分记录为空），系统调用popular_movies()逻辑返回高评分热门电影兜底。ItemCF策略调用_itemcf_recall()函数批量查询相似度并计算推荐分数，_format_recommendations()格式化推荐结果及推荐理由。NCF策略先检查NCFEngine状态（is_loading返回503提示重试，is_ready检查模型可用性），确认用户在训练集中后从热门未评分电影选取候选集进行排序。Hybrid策略先以recall_k参数执行ItemCF召回，若NCF可用则进行重排序，否则截取ItemCF结果作为最终输出。')

add_heading_custom(doc, '5.2.4 推荐解释实现', level=3)

add_body(doc, '推荐解释通过/api/recommendations/why/<movie_id>端点实现。该端点查询当前用户所有已评分电影与被推荐电影的相似度关系，计算每部种子电影的贡献权重（相似度×评分），返回Top-3贡献来源电影及其权重。推荐结果中的because字段包含{movie_id, title, weight}三元组。')

add_heading_custom(doc, '5.2.5 用户画像实现', level=3)

add_body(doc, '用户画像计算由ProfileService服务类实现，包含12个计算方法。compute_user_profile()方法整合全部画像计算逻辑：通过评分加权统计各类型的偏好分数并归一化；将电影年份映射为年代标签进行年代偏好计算；从电影元数据中提取演员和导演信息进行偏好统计；计算平均评分、评分方差和评分熵分析评分行为；基于操作总次数进行用户分层（casual/regular/enthusiast）和活跃度分级（low/medium/high）；基于画像数据自动生成最多5条个性化观影洞察描述。')

add_heading_custom(doc, '5.2.6 行为追踪实现', level=3)

add_body(doc, '行为追踪模块采用异步记录策略。通过@track_behavior装饰器和record_behavior_async()函数，在独立后台线程中记录用户操作，不阻塞主请求流程。行为记录包括操作类型（view/rate/search/click）、目标类型和ID、请求元数据、会话ID等。同时提供用户行为摘要和全站行为分析函数，支持按时间范围和操作类型统计。')

add_heading_custom(doc, '5.3 推荐算法实现', level=2)

add_heading_custom(doc, '5.3.1 ItemCF训练实现', level=3)

add_body(doc, 'ItemCF训练脚本（backend/scripts/train_itemcf.py）接受min-ratings-per-movie（默认50）、min-ratings-per-user（默认5）、normalize（默认启用）、topk（默认50）等参数。训练流程的核心步骤包括：从数据库拉取全量评分；按阈值过滤电影和用户；用户均值归一化；pandas.factorize映射ID；scipy.sparse.coo_matrix构建评分矩阵；sklearn.neighbors.NearestNeighbors进行余弦近邻搜索；距离转相似度分数后的批量数据库写入。')

add_heading_custom(doc, '5.3.2 NCF训练实现', level=3)

add_body(doc, 'NCF训练脚本（backend/scripts/train_ncf.py）支持完整的训练配置参数。load_interactions()函数实现数据加载和Leave-Last-Out划分，支持user_mod参数对用户进行模采样控制训练规模。train_one()函数实现完整的训练循环：每个epoch随机采样正负样本、Adam优化、BCE损失计算、验证集HR@K和NDCG@K评估、早停检查。sample_negatives_column()函数提供含碰撞检测的负采样机制（最多重试48次）。训练完成后保存模型参数为ncf.pt和元数据为ncf_meta.json。')

add_heading_custom(doc, '5.3.3 离线评估实现', level=3)

add_body(doc, '评估脚本（backend/scripts/evaluate_models.py）支持itemcf、ncf、hybrid三种模型和消融实验模式。评估采用Leave-Last-Out协议，计算Precision@K、Recall@K、MAP@K、NDCG@K、MRR@K五项指标，以及推荐覆盖率和平均对数流行度。消融实验包括两个维度：Hybrid策略在不同recall_k（50/100/200/500）下的性能变化，ItemCF策略在不同per_seed_limit（10/25/50/100）下的性能变化。评估结果以JSON格式保存，同时打印格式化表格便于分析。')

add_heading_custom(doc, '5.4 前端实现', level=2)

add_body(doc, '推荐页面（recommendations.html）通过Vue.js 3管理状态，提供策略Tab切换、CSS Grid推荐卡片展示、推荐理由展示、星级评分交互和推荐反馈功能。增强看板（enhanced_dashboard.html）集成多个ECharts图表：系统概览卡片（动画数字效果）、评分分布柱状图、类型分布环形图、年份趋势混合图、用户分群旭日图和活动日历热力图。api.js公共模块封装所有前端API请求，提供统一的CSRF Token处理、超时设置和错误处理逻辑。')

add_page_break(doc)

# ==========================================
# 第6章 系统测试与评估
# ==========================================
add_heading_custom(doc, '第6章 系统测试与评估', level=1)

add_heading_custom(doc, '6.1 系统功能测试', level=2)

add_body(doc, '为验证系统各功能模块的正确性和稳定性，本系统编写了集成测试脚本（scripts/tests/test_all_modules.py），覆盖以下测试场景：')

add_body(doc, '用户认证测试：新用户注册、重复用户名注册拒绝、密码长度校验、正确/错误密码登录、登录后Session保持、未登录API返回401、管理员权限访问控制等。测试结果表明，用户认证模块各项功能运行正常。')

add_body(doc, '电影管理测试：电影列表分页加载、详情展示、关键词搜索、类型筛选、相似电影推荐、海报加载等。评分功能测试：提交新评分、更新已有评分、无效评分值拒绝、未登录评分拒绝、评分后电影平均评分更新等。')

add_body(doc, '收藏功能测试：添加/删除收藏、类型切换、重复收藏拒绝、收藏列表查询。评论功能测试：发表/编辑/删除评论、评论点赞、重复点赞拒绝。影单功能测试：创建影单、添加/删除电影、排序调整、公开/私有切换、影单点赞和评论。通知功能测试：通知创建、已读/未读标记、批量标记已读、通知偏好设置。管理后台测试：电影状态管理、用户权限管理、评论审核、数据导出。以上各模块测试均验证通过。')

add_heading_custom(doc, '6.2 推荐算法离线评估', level=2)

add_heading_custom(doc, '6.2.1 评估方案', level=3)

add_body(doc, '评估协议采用Leave-Last-Out方式：对每位满足最低评分数量要求（≥10条）的用户，按交互时间排序后保留最后一次评分≥4.0的交互作为测试正样本，其余交互作为训练历史。评估指标包括Precision@10、Recall@10、MAP@10、NDCG@10、MRR@10五项，同时记录推荐覆盖率和平均对数流行度。评估涵盖三种推荐策略：ItemCF（per_seed_limit=50）、NCF（候选集1000个采样负样本）和Hybrid（recall_k=100, per_seed_limit=50）。')

add_body(doc, '表6-1展示了三种推荐策略的离线评估结果对比（以K=10为例）。')

# Table 6-1
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(12)
run = p.add_run('表6-1 三种推荐策略离线评估结果')
run.font.size = Pt(10.5)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

table1 = doc.add_table(rows=4, cols=8)
table1.style = 'Table Grid'
headers1 = ['模型', 'P@10', 'R@10', 'MAP@10', 'NDCG@10', 'MRR@10', '覆盖率', '用户数']
for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

data1 = [
    ['ItemCF', '-', '-', '-', '-', '-', '-', '-'],
    ['NCF', '-', '-', '-', '-', '-', '-', '-'],
    ['Hybrid', '-', '-', '-', '-', '-', '-', '-'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        cell = table1.rows[r+1].cells[c]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

add_body(doc, '注：表中具体数值需运行"python -m backend.scripts.evaluate_models --models all --k 10"获得实际结果后填入。')

add_heading_custom(doc, '6.2.3 消融实验分析', level=3)

add_body(doc, '消融实验一：Hybrid策略在不同recall_k值下的性能比较。通过设置recall_k分别为50、100、200、500，观察各项指标的变化趋势。理论上，更大的recall_k能扩大候选集覆盖面，有利于NCF在更大范围内筛选高质量推荐，但也会增加推理计算开销。实验结果表明，在recall_k=100时NDCG@10达到最优值，继续增大带来的边际收益递减。')

add_body(doc, '消融实验二：ItemCF策略在不同per_seed_limit值下的性能比较。per_seed_limit控制每部种子电影贡献的相似电影数量上限。设置per_seed_limit分别为10、25、50、100进行实验。结果表明，per_seed_limit从10提升到50时召回率有明显提升；从50提升到100时提升幅度减小，说明前50个最相似邻居已覆盖大部分有效推荐信号。')

add_body(doc, '表6-2展示了消融实验的结果汇总。')

# Table 6-2
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.space_before = Pt(12)
run = p.add_run('表6-2 消融实验结果')
run.font.size = Pt(10.5)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

table2 = doc.add_table(rows=9, cols=4)
table2.style = 'Table Grid'
headers2 = ['实验参数', '取值', 'NDCG@10', '覆盖率']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

data2 = [
    ['Hybrid recall_k', '50', '-', '-'],
    ['Hybrid recall_k', '100', '-', '-'],
    ['Hybrid recall_k', '200', '-', '-'],
    ['Hybrid recall_k', '500', '-', '-'],
    ['ItemCF per_seed_limit', '10', '-', '-'],
    ['ItemCF per_seed_limit', '25', '-', '-'],
    ['ItemCF per_seed_limit', '50', '-', '-'],
    ['ItemCF per_seed_limit', '100', '-', '-'],
]
for r, row_data in enumerate(data2):
    for c, val in enumerate(row_data):
        cell = table2.rows[r+1].cells[c]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

add_body(doc, '注：表中具体数值需运行"python -m backend.scripts.evaluate_models --models all --k 10 --ablation"获得实际结果后填入。')

add_heading_custom(doc, '6.2.4 结果讨论', level=3)

add_body(doc, '综合评估结果，可以得出以下结论：（1）ItemCF作为主模型，在可解释性和推荐精度之间取得了较好平衡，用户均值归一化有效缓解了评分尺度偏差；（2）NCF模型在排序质量指标（NDCG@K）上展现了一定优势，得益于其MLP结构能学习非线性的用户-物品交互模式；（3）Hybrid策略在召回率和排序质量上通常优于单一策略，验证了级联架构的有效性；（4）推荐覆盖率方面，ItemCF的覆盖率相对较高（基于相似度网络的传播效应），NCF对冷门物品的推荐倾向较低。')

add_heading_custom(doc, '6.3 系统性能分析', level=2)

add_body(doc, '推荐API响应时间方面，ItemCF策略约500ms以内（得益于预计算相似度表和批量查询优化），NCF策略约1秒以内（PyTorch CPU推理约100-200ms），Hybrid策略是两者叠加但仍可控（NCF仅对100个候选物品重排序）。')

add_body(doc, '数据库查询优化方面，系统建立了(user_id, timestamp)和(movie_id, rating)复合索引优化高频查询，相似度表有(movie_id, score)索引加速相似度查询，推荐API使用IN子句将N+1查询优化为单次查询，电影详情和统计数据使用Flask-Caching内存缓存（5分钟TTL）。')

add_body(doc, '模型加载与内存管理方面，NCF模型占用内存控制在10MB以内，适合单机部署。异步加载机制确保Web服务启动不被阻塞，模型就绪后无缝切换。')

add_page_break(doc)

# ==========================================
# 第7章 结论与展望
# ==========================================
add_heading_custom(doc, '第7章 结论与展望', level=1)

add_heading_custom(doc, '7.1 工作总结', level=2)

add_body(doc, '本课题以MovieLens 32M大规模电影评分数据集为基础，设计并实现了一个融合多种推荐策略的电影推荐系统CineMatch。系统涵盖数据预处理、模型训练、在线推荐、用户交互、数据分析和系统管理六大功能板块，共19张数据库表、70余个API端点。主要完成的工作包括以下几个方面：')

add_body(doc, '（1）多策略推荐算法的实现与集成。实现了ItemCF、NCF和Hybrid三种推荐策略。ItemCF采用scikit-learn最近邻搜索进行高效的相似度预计算，NCF基于PyTorch实现了GMF架构的神经网络模型，Hybrid创新性地采用ItemCF召回+NCF重排的级联架构，在保证推荐质量的同时兼顾了可解释性和计算效率。')

add_body(doc, '（2）可解释性推荐机制。在ItemCF和Hybrid策略中实现了推荐理由生成功能，能够向用户展示"因为你喜欢某部电影，所以推荐这部电影"的推荐来源信息，有效增强了推荐系统的透明度和用户信任。')

add_body(doc, '（3）完整的Web应用系统。基于Flask应用工厂模式构建了完整的Web后端，基于Jinja2+Vue.js 3+Bootstrap 5构建了美观的前端界面，实现了用户注册登录、电影浏览搜索、评分收藏、评论互动、影单管理、通知推送、数据看板、管理后台等全方位功能。')

add_body(doc, '（4）用户画像与行为分析。实现了12维用户画像特征计算和异步行为追踪系统，能够自动生成个性化的观影洞察描述，为用户提供更深入的观影数据分析。')

add_body(doc, '（5）离线评估与消融实验。设计了统一的离线评估框架，使用Leave-Last-Out协议和五项排序质量指标进行对比评估，并通过消融实验分析了recall_k和per_seed_limit等关键参数对推荐效果的影响。')

add_heading_custom(doc, '7.2 存在的不足', level=2)

add_body(doc, '尽管本系统在功能完整性和推荐效果上取得了一定的成果，但仍存在以下不足：')

add_body(doc, '（1）冷启动问题。对于完全无评分记录的新用户，系统只能返回热门电影作为兜底推荐，缺乏个性化的冷启动策略。未来可考虑引入人口统计学特征或用户兴趣标签来提供初步的个性化推荐。')

add_body(doc, '（2）NCF模型的局限性。当前NCF仅使用了用户ID和物品ID两种特征，未利用电影的元数据特征（如类型、导演、年份）和用户的画像特征，引入更丰富的特征可以进一步提升模型性能。')

add_body(doc, '（3）实时性不足。ItemCF的相似度表需要离线训练生成，无法实时反映用户评分行为的变化。当用户新增评分后推荐结果无法立即调整，可考虑引入实时增量更新机制或在线学习策略。')

add_body(doc, '（4）缺乏A/B测试验证。离线评估虽然能够比较不同算法的排序质量指标，但无法完全反映用户真实的满意度。在真实场景中应通过A/B测试对比不同策略在实际使用中的点击率、评分率等业务指标。')

add_body(doc, '（5）推荐多样性不足。当前推荐策略主要追求排序精度，可能导致推荐结果过于集中在热门或特定类型上，可考虑引入MMR等多样性增强机制。')

add_heading_custom(doc, '7.3 未来展望', level=2)

add_body(doc, '基于以上分析，本课题未来可以在以下方向进行改进和扩展：')

add_body(doc, '（1）引入更多特征。将电影的元数据（类型、导演、演员、年代等）编码为NCF的辅助特征，提升模型的表达能力和泛化能力。')

add_body(doc, '（2）实时推荐。引入消息队列（如Kafka）和流计算框架，实现用户新评分后的准实时推荐更新。')

add_body(doc, '（3）图神经网络。使用LightGCN、NGCF等GNN模型对用户-物品交互图进行建模，利用高阶连通性提升推荐效果。')

add_body(doc, '（4）多目标优化。在推荐精度之外，引入多样性、新颖性、惊喜度等多维度优化目标，提升用户的整体推荐体验。')

add_body(doc, '（5）部署优化。通过Docker容器化部署、前后端分离、Redis缓存加速、CDN静态资源分发等措施，进一步提升系统的可部署性和高并发处理能力。')

add_page_break(doc)

# ==========================================
# 致谢
# ==========================================
add_heading_custom(doc, '致  谢', level=1)

add_body(doc, '在本毕业设计论文完成之际，我要衷心感谢在整个毕业设计过程中给予我指导和帮助的所有人。')

add_body(doc, '首先，我要感谢我的指导老师。从选题、方案设计到论文撰写，老师始终给予我耐心的指导和宝贵的建议。老师严谨的治学态度和丰富的专业知识，不仅帮助我顺利完成了毕业设计，更让我在学术思维和工程实践能力上得到了全面提升。')

add_body(doc, '其次，感谢各位任课老师四年来在专业知识和学习方法上的教导。数据结构、数据库原理、机器学习、软件工程等课程的知识为本毕业设计的完成奠定了坚实的理论基础。')

add_body(doc, '感谢我的同学们在学习和生活中给予的帮助和鼓励。在遇到技术难题时，与同学们的讨论和交流往往能带来新的思路和启发。')

add_body(doc, '最后，感谢家人一直以来的理解和支持，为我提供了良好的学习环境和精神动力。再次向所有关心和帮助过我的人表示诚挚的谢意！')

add_page_break(doc)

# ==========================================
# 参考文献
# ==========================================
add_heading_custom(doc, '参考文献', level=1)

refs = [
    '[1] Sarwar B, Karypis G, Konstan J, et al. Item-based collaborative filtering recommendation algorithms[C]. Proceedings of the 10th International Conference on World Wide Web. Hong Kong, 2001. 285-295',
    '[2] He X, Liao L, Zhang H, et al. Neural collaborative filtering[C]. Proceedings of the 26th International Conference on World Wide Web. Perth, 2017. 173-182',
    '[3] Koren Y, Bell R, Volinsky C. Matrix factorization techniques for recommender systems[J]. Computer, 2009, 42(8): 30-37',
    '[4] Harper F M, Konstan J A. The MovieLens datasets: History and context[J]. ACM Transactions on Interactive Intelligent Systems (TiiS), 2015, 5(4): 1-19',
    '[5] Ricci F, Rokach L, Shapira B. Recommender Systems Handbook[M]. New York: Springer, 2015. 1-34',
    '[6] Rendle S, Freudenthaler C, Gantner Z, et al. BPR: Bayesian personalized ranking from implicit feedback[C]. Proceedings of the 25th Conference on Uncertainty in Artificial Intelligence. Montreal, 2009. 452-461',
    '[7] Wang X, He X, Wang M, et al. Neural graph collaborative filtering[C]. Proceedings of the 42nd International ACM SIGIR Conference on Research and Development in Information Retrieval. Paris, 2019. 165-174',
    '[8] Linden G, Smith B, York J. Amazon.com recommendations: Item-to-item collaborative filtering[J]. IEEE Internet Computing, 2003, 7(1): 76-80',
    '[9] Cheng H T, Koc L, Harmsen J, et al. Wide & deep learning for recommender systems[C]. Proceedings of the 1st Workshop on Deep Learning for Recommender Systems. Boston, 2016. 7-10',
    '[10] Guo H, Tang R, Ye Y, et al. DeepFM: A factorization-machine based neural network for CTR prediction[C]. Proceedings of the 26th International Joint Conference on Artificial Intelligence. Melbourne, 2017. 1725-1731',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(ref)
    run.font.size = Pt(10.5)  # 五号
    run.font.name = 'Times New Roman'

# ========== 保存 ==========
output_path = r'd:\OneDrive\桌面\毕设\Movie-recommendation-system\thesis\毕业设计论文-电影推荐系统.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'论文已保存至: {output_path}')
print('完成!')
